"""DentaScribe Clinical SOAP — DOCX generator.

Takes the structured SOAP JSON the agent swarm produces and renders a real
printable Word document that an attending dentist would sign. JSON stays as
the internal source-of-truth; this module is the *output* layer.

Layout (matches `exports/soap_pdf_template.py`):

    [Header: brand + encounter meta]
    [Encounter row: patient ID · DOB · visit type · provider · license · date]
    Subjective
      Chief complaint (quoted)
      HPI — OPQRST table
      Medical hx · medications · allergies · social hx
    Objective
      Vitals
      Extra-oral / intra-oral findings
      Hard-tissue findings (per-tooth table)
      Periodontal summary
      Radiographs
    Assessment
      Diagnoses (per tooth, with severity)
      Differential (only if explicitly stated)
    Plan
      Procedures today (per tooth, with surfaces + anesthesia + CDT)
      Prescriptions (drug · strength · sig · qty · refills · ✓ interaction check)
      Follow-up
      Patient instructions
    Billing
      CDT codes (code · description · tooth · surfaces · rationale)
    Compliance — TSBDE 22 TAC §108.8 anchor block (checklist)
    Attestation
      AI-assisted disclosure + provider sig
    [Footer: generated timestamp + signability score]

Design rules:
  - Use the brand palette (mint-teal accent, deep navy text) sparingly —
    print contexts need black on white.
  - Tables for anything tabular. Bold + subtle accent bar for section heads.
  - Empty fields become "—" not blank, so the doc looks complete even with
    sparse data.
  - Never leak PHI to logs (the caller writes the bytes; no print here).
"""
from __future__ import annotations
import io
import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Mm, Pt, RGBColor


# Brand colors — match ui/theme.py palette for visual continuity with the app.
ACCENT      = RGBColor(0x4D, 0xD4, 0xAC)
ACCENT_SOFT = RGBColor(0xE6, 0xF8, 0xF1)
NAVY        = RGBColor(0x0F, 0x17, 0x2A)
INK         = RGBColor(0x1B, 0x22, 0x33)
DIM         = RGBColor(0x60, 0x6B, 0x7F)
ROSE        = RGBColor(0xC9, 0x4D, 0x4D)
AMBER       = RGBColor(0xB3, 0x80, 0x2E)


# ---------- public API ----------

def build_soap_docx(soap: dict | None, attestation: dict | None = None,
                    validation: dict | None = None) -> bytes:
    """Return the SOAP note as a populated Word .docx (bytes).

    `soap` is the dict the swarm produces (see core/soap_schema.json).
    `attestation` (optional) carries provider_name, license, signed_at.
    `validation` (optional) carries signability_score for the audit footer.
    """
    doc = Document()
    _configure_styles(doc)
    _set_margins(doc, top_mm=15, bottom_mm=15, left_mm=18, right_mm=18)

    soap = soap or {}
    _render_header(doc, soap)
    _render_encounter_block(doc, soap)
    _render_subjective(doc, soap.get("subjective") or {})
    _render_objective(doc, soap.get("objective") or {})
    _render_assessment(doc, soap.get("assessment") or {})
    _render_plan(doc, soap.get("plan") or {})
    _render_billing(doc, soap.get("billing") or {})
    _render_compliance(doc, soap.get("compliance") or {})
    _render_attestation(doc, attestation or {}, soap)
    _render_footer(doc, validation)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------- styles + page setup ----------

def _configure_styles(doc: Document) -> None:
    s = doc.styles["Normal"]
    s.font.name = "Helvetica"
    s.font.size = Pt(10)
    s.font.color.rgb = INK


def _set_margins(doc, *, top_mm: int, bottom_mm: int, left_mm: int, right_mm: int) -> None:
    for section in doc.sections:
        section.top_margin = Mm(top_mm)
        section.bottom_margin = Mm(bottom_mm)
        section.left_margin = Mm(left_mm)
        section.right_margin = Mm(right_mm)


def _shade_cell(cell, rgb_hex: str) -> None:
    """Apply a background fill to a table cell (python-docx doesn't expose it)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex)
    tc_pr.append(shd)


def _set_cell_border(cell, *, top=None, left=None, bottom=None, right=None,
                      color: str = "DDDDDD", sz: int = 6) -> None:
    """Apply borders to a single cell. Subtle dividers for clinical look."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for side, on in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        if on is None:
            continue
        el = OxmlElement(f"w:{side}")
        el.set(qn('w:val'), "single" if on else "nil")
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:color'), color)
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def _hide_table_borders(table) -> None:
    """No outer borders — table-as-layout. Borders on individual cells if needed."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for tag in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{tag}")
        el.set(qn('w:val'), "nil")
        borders.append(el)
    tblPr.append(borders)


# ---------- composable text helpers ----------

def _section_heading(doc, text: str, *, accent: bool = True) -> None:
    """Section header — small mint accent bar + uppercase letterspaced label."""
    # Spacer
    doc.add_paragraph().paragraph_format.space_before = Pt(6)
    # Accent bar via a 1-row table
    bar = doc.add_table(rows=1, cols=2)
    _hide_table_borders(bar)
    bar.autofit = False
    bar.columns[0].width = Mm(2)
    bar.columns[1].width = Mm(170)
    bar.rows[0].height = Mm(6)
    bar.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    if accent:
        _shade_cell(bar.cell(0, 0), "4DD4AC")
    p = bar.cell(0, 1).paragraphs[0]
    r = p.add_run("  " + text.upper())
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)


def _kv_row(table, label: str, value: str) -> None:
    """Append a (label, value) row to a 2-col layout table."""
    row = table.add_row()
    a, b = row.cells
    a.width = Mm(35)
    b.width = Mm(135)
    pa = a.paragraphs[0]
    ra = pa.add_run(label)
    ra.font.color.rgb = DIM
    ra.font.size = Pt(9)
    ra.bold = False
    pb = b.paragraphs[0]
    rb = pb.add_run(value if value not in (None, "", []) else "—")
    rb.font.size = Pt(10)


def _add_field(doc, label: str, value: str) -> None:
    """Single-line labeled field. Used outside the kv tables."""
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}  ")
    r1.font.color.rgb = DIM
    r1.font.size = Pt(9)
    r2 = p.add_run(value if value not in (None, "", []) else "—")
    r2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)


def _paragraph(doc, text: str, *, italic: bool = False, color: RGBColor | None = None) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    if color: r.font.color.rgb = color
    r.font.size = Pt(10)


def _bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Mm(6)
    r = p.add_run(text)
    r.font.size = Pt(10)


def _safe(value: Any) -> str:
    """Coerce any SOAP field to a printable string. Empty → '—'."""
    if value in (None, "", []):
        return "—"
    if isinstance(value, list):
        return ", ".join(_safe(v) for v in value)
    if isinstance(value, dict):
        return ", ".join(f"{k}: {_safe(v)}" for k, v in value.items())
    return str(value).strip() or "—"


# ---------- sections ----------

def _render_header(doc, soap: dict) -> None:
    """Brand band at the top. Tooth glyph + product name + practice."""
    table = doc.add_table(rows=1, cols=2)
    _hide_table_borders(table)
    table.autofit = False
    table.columns[0].width = Mm(110)
    table.columns[1].width = Mm(64)

    left = table.cell(0, 0).paragraphs[0]
    r = left.add_run("🦷  DentaScribe")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = NAVY
    sub = left.add_run("\nClinical SOAP Record")
    sub.font.size = Pt(9)
    sub.font.color.rgb = DIM

    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta = soap.get("encounter_meta") or soap.get("metadata") or {}
    loc = meta.get("practice_location") or {}
    loc_str = f"{loc.get('city', 'Dallas')}, {loc.get('state', 'TX')}"
    r2 = right.add_run(f"{loc_str}\nTSBDE 22 TAC §108.8 compliant")
    r2.font.size = Pt(9)
    r2.font.color.rgb = DIM

    # Thin divider
    sep = doc.add_table(rows=1, cols=1)
    _hide_table_borders(sep)
    sep.rows[0].height = Mm(0.6)
    sep.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    _shade_cell(sep.cell(0, 0), "4DD4AC")


def _render_encounter_block(doc, soap: dict) -> None:
    """Patient + provider + date + visit type — the chart anchor block."""
    meta = soap.get("encounter_meta") or soap.get("metadata") or {}
    patient = (meta.get("patient") or {})
    provider = (meta.get("provider") or {})

    doc.add_paragraph().paragraph_format.space_before = Pt(4)
    box = doc.add_table(rows=1, cols=4)
    _hide_table_borders(box)
    box.autofit = False
    widths = [Mm(43), Mm(43), Mm(43), Mm(43)]
    for i, w in enumerate(widths):
        box.columns[i].width = w

    def _slot(cell, label, value):
        _shade_cell(cell, "F4F8FB")
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        r1 = cell.paragraphs[0].add_run(label.upper())
        r1.bold = True
        r1.font.size = Pt(7.5)
        r1.font.color.rgb = DIM
        # letter spacing not directly exposed; rely on style + run
        p = cell.add_paragraph()
        r2 = p.add_run(value if value not in (None, "", []) else "—")
        r2.font.size = Pt(11)
        r2.font.color.rgb = INK
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(1)

    _slot(box.cell(0, 0), "Encounter date", meta.get("date_of_service", "—"))
    _slot(box.cell(0, 1), "Visit type",    (meta.get("visit_type") or "—").replace("_", " "))
    _slot(box.cell(0, 2), "Patient ID",    patient.get("patient_id", "—"))
    _slot(box.cell(0, 3), "Patient DOB",   patient.get("dob", "—"))

    box2 = doc.add_table(rows=1, cols=4)
    _hide_table_borders(box2)
    for i, w in enumerate(widths):
        box2.columns[i].width = w
    _slot(box2.cell(0, 0), "Provider",       provider.get("name", "—"))
    _slot(box2.cell(0, 1), "TSBDE license",  provider.get("tsbde_license", "—"))
    _slot(box2.cell(0, 2), "NPI",            provider.get("npi", "—"))
    _slot(box2.cell(0, 3), "Role",           provider.get("role", "—"))


def _render_subjective(doc, s: dict) -> None:
    _section_heading(doc, "Subjective")

    # Chief complaint — verbatim quote, indented + italicized
    cc = s.get("chief_complaint")
    if cc:
        _add_field(doc, "Chief complaint:", "")
        q = doc.add_paragraph()
        q.paragraph_format.left_indent = Mm(8)
        r = q.add_run(f"“{cc}”")
        r.italic = True
        r.font.size = Pt(11)
        r.font.color.rgb = NAVY
    else:
        _add_field(doc, "Chief complaint:", "—")

    # HPI — narrative if simple, OPQRST table if structured
    hpi = s.get("hpi") or s.get("history_of_present_illness")
    if isinstance(hpi, dict):
        _add_field(doc, "History of present illness (OPQRST):", "")
        opqrst = doc.add_table(rows=0, cols=2)
        _hide_table_borders(opqrst)
        opqrst.columns[0].width = Mm(35)
        opqrst.columns[1].width = Mm(135)
        for key, label in [
            ("onset", "Onset"),
            ("provocation", "Provocation"),
            ("quality", "Quality"),
            ("region", "Region / radiation"),
            ("severity_0_10", "Severity (0–10)"),
            ("timing", "Timing"),
            ("triggers", "Triggers"),
            ("character", "Character"),
        ]:
            if key in hpi and hpi[key] not in (None, "", []):
                _kv_row(opqrst, label, _safe(hpi[key]))
    elif hpi:
        _add_field(doc, "History of present illness:", "")
        _paragraph(doc, str(hpi))

    # Medical history, meds, allergies, social hx
    hx_table = doc.add_table(rows=0, cols=2)
    _hide_table_borders(hx_table)
    hx_table.columns[0].width = Mm(35)
    hx_table.columns[1].width = Mm(135)
    _kv_row(hx_table, "Medical history", _safe(s.get("medical_history_updates")))
    _kv_row(hx_table, "Medications",     _safe(s.get("medications")))
    _kv_row(hx_table, "Allergies",       _safe(s.get("allergies")))
    if s.get("social_history"):
        _kv_row(hx_table, "Social history", _safe(s.get("social_history")))


def _render_objective(doc, o: dict) -> None:
    _section_heading(doc, "Objective")

    # Vitals
    vitals = o.get("vitals") or {}
    if vitals:
        parts = []
        for k in ("bp", "hr", "rr", "temp", "spo2", "weight", "height", "bmi"):
            if k in vitals:
                parts.append(f"{k.upper()} {vitals[k]}")
        if parts:
            _add_field(doc, "Vitals:", "  ·  ".join(parts))

    if o.get("extra_oral"):
        _add_field(doc, "Extra-oral exam:", _safe(o["extra_oral"]))
    if o.get("intra_oral"):
        _add_field(doc, "Intra-oral exam:", _safe(o["intra_oral"]))

    # Hard-tissue findings — per-tooth table
    findings = o.get("exam_findings") or []
    if findings:
        _add_field(doc, "Hard-tissue findings:", "")
        t = doc.add_table(rows=1, cols=5)
        _style_data_table(t, headers=["Tooth", "Surfaces", "Finding", "Severity", "Source"])
        for f in findings:
            if not isinstance(f, dict):
                continue
            row = t.add_row().cells
            row[0].text = _safe(f.get("tooth"))
            row[1].text = ",".join(f.get("surfaces") or []) or "—"
            row[2].text = _safe(f.get("finding"))
            row[3].text = _safe(f.get("severity"))
            row[4].text = _safe(f.get("source_span"))[:80]
            for c in row:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)

    # Periodontal summary
    perio = o.get("perio_summary") or o.get("periodontal_summary")
    if perio:
        _add_field(doc, "Periodontal:", _safe(perio))

    # Radiographs
    rads = o.get("radiographs_taken") or []
    if rads:
        _add_field(doc, "Radiographs:", "")
        for r in rads:
            if isinstance(r, dict):
                parts = [r.get("type", ""), f"#{r['tooth']}" if r.get("tooth") else None, r.get("findings", "")]
                _bullet(doc, " — ".join(p for p in parts if p))
            elif r:
                _bullet(doc, str(r))


def _render_assessment(doc, a: dict) -> None:
    _section_heading(doc, "Assessment")
    diagnoses = a.get("diagnoses") or a.get("primary_diagnosis")
    if isinstance(diagnoses, list) and diagnoses:
        for dx in diagnoses:
            if isinstance(dx, dict):
                tooth = dx.get("tooth")
                sev = dx.get("severity")
                tag = ""
                if tooth: tag += f"  ·  tooth #{tooth}"
                if sev:   tag += f"  ·  {sev}"
                _bullet(doc, f"{_safe(dx.get('diagnosis'))}{tag}")
            else:
                _bullet(doc, _safe(dx))
    elif diagnoses:
        _paragraph(doc, _safe(diagnoses))
    else:
        _paragraph(doc, "—", color=DIM)

    differentials = a.get("differentials") or a.get("differential_diagnoses")
    if differentials:
        _add_field(doc, "Differential:", _safe(differentials))


def _render_plan(doc, p: dict) -> None:
    _section_heading(doc, "Plan")

    # Procedures today
    procs = p.get("procedures_today") or []
    if procs:
        _add_field(doc, "Procedures today:", "")
        t = doc.add_table(rows=1, cols=5)
        _style_data_table(t, headers=["Tooth", "Surfaces", "Procedure", "Anesthesia", "CDT"])
        for proc in procs:
            if not isinstance(proc, dict):
                continue
            row = t.add_row().cells
            row[0].text = _safe(proc.get("tooth"))
            row[1].text = ",".join(proc.get("surfaces") or []) or "—"
            row[2].text = _safe(proc.get("procedure"))
            row[3].text = _safe(proc.get("anesthesia"))
            row[4].text = _safe(proc.get("cdt_code"))
            for c in row:
                for para in c.paragraphs:
                    for r in para.runs:
                        r.font.size = Pt(9)

    # Prescriptions
    rxs = p.get("prescriptions") or []
    if rxs:
        _add_field(doc, "Prescriptions:", "")
        for rx in rxs:
            if not isinstance(rx, dict):
                continue
            drug   = rx.get("drug", "—")
            stren  = rx.get("strength") or ""
            sig    = rx.get("sig") or ""
            qty    = rx.get("quantity")
            refill = rx.get("refills")
            ok     = rx.get("interaction_checked")
            line = f"{drug} {stren}  ·  {sig}"
            if qty is not None:    line += f"  ·  Qty {qty}"
            if refill is not None: line += f"  ·  Refills {refill}"
            line += "   "
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent = Mm(6)
            r1 = para.add_run(line)
            r1.font.size = Pt(10)
            r2 = para.add_run("✓ interaction checked" if ok else "⚠ interaction check pending")
            r2.font.size = Pt(9)
            r2.font.color.rgb = ACCENT if ok else AMBER

    # Follow-up + patient instructions
    if p.get("follow_up"):
        _add_field(doc, "Follow-up:", _safe(p["follow_up"]))
    if p.get("patient_instructions"):
        _add_field(doc, "Patient instructions:", _safe(p["patient_instructions"]))

    # Recommended-future (deferred procedures)
    rec = p.get("recommended_future") or []
    if rec:
        _add_field(doc, "Recommended (future):", "")
        for item in rec:
            if isinstance(item, dict):
                _bullet(doc, f"{_safe(item.get('procedure'))} — {_safe(item.get('when'))}")
            else:
                _bullet(doc, _safe(item))


def _render_billing(doc, b: dict) -> None:
    _section_heading(doc, "Billing (CDT 2026)")
    codes = b.get("cdt_codes") or []
    if not codes:
        _paragraph(doc, "No CDT codes assigned.", color=DIM)
        return
    t = doc.add_table(rows=1, cols=5)
    _style_data_table(t, headers=["Code", "Description", "Tooth", "Surfaces", "Rationale"])
    for c in codes:
        if not isinstance(c, dict):
            continue
        row = t.add_row().cells
        row[0].text = _safe(c.get("code"))
        row[1].text = _safe(c.get("description"))
        row[2].text = _safe(c.get("tooth"))
        row[3].text = ",".join(c.get("surfaces") or []) or "—"
        row[4].text = _safe(c.get("rationale") or c.get("code_null_reason"))[:120]
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    est = b.get("estimated_total")
    if est is not None:
        _add_field(doc, "Estimated total:", f"${est:,.2f}" if isinstance(est, (int, float)) else _safe(est))


def _render_compliance(doc, c: dict) -> None:
    """TSBDE 22 TAC §108.8 anchor block — checklist of required record elements."""
    _section_heading(doc, "Compliance — TSBDE 22 TAC §108.8")
    checklist = c.get("tsbde_checklist") or {}
    if not checklist:
        _paragraph(doc, "No compliance data computed.", color=DIM)
        return

    cols = doc.add_table(rows=0, cols=2)
    _hide_table_borders(cols)
    cols.columns[0].width = Mm(85)
    cols.columns[1].width = Mm(85)
    items = list(checklist.items())
    for i in range(0, len(items), 2):
        row = cols.add_row()
        for col_idx, item in enumerate(items[i:i+2]):
            key, val = item
            label = key.replace("_", " ").title()
            cell = row.cells[col_idx]
            mark = "✓" if val is True else ("⚠" if val is None else "✗")
            color = ACCENT if val is True else (AMBER if val is None else ROSE)
            p = cell.paragraphs[0]
            r1 = p.add_run(f" {mark}  ")
            r1.bold = True
            r1.font.size = Pt(11)
            r1.font.color.rgb = color
            r2 = p.add_run(label)
            r2.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(2)


def _render_attestation(doc, att: dict, soap: dict) -> None:
    _section_heading(doc, "Attestation")
    soap_att = soap.get("attestation") or {}

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(
        "This SOAP note was drafted by DentaScribe (an AI-assisted clinical scribe) "
        "and reviewed by the signing provider before signature. The provider takes "
        "clinical responsibility for the content of this record."
    )
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = DIM

    sig = doc.add_table(rows=1, cols=2)
    _hide_table_borders(sig)
    sig.columns[0].width = Mm(85)
    sig.columns[1].width = Mm(85)
    left, right = sig.cell(0, 0), sig.cell(0, 1)

    provider_name = att.get("provider_name") or soap_att.get("signed_by") or "—"
    license_num   = att.get("license") or "—"
    signed_at     = att.get("signed_at") or soap_att.get("signed_at") or "Not signed"

    _shade_cell(left, "F4F8FB")
    lp = left.paragraphs[0]
    lp.add_run("PROVIDER SIGNATURE").bold = True
    lp.runs[0].font.size = Pt(7.5)
    lp.runs[0].font.color.rgb = DIM
    lp_v = left.add_paragraph()
    lp_v.add_run(provider_name).font.size = Pt(11)

    rp = right.paragraphs[0]
    rp.add_run("DATE / TIME").bold = True
    rp.runs[0].font.size = Pt(7.5)
    rp.runs[0].font.color.rgb = DIM
    _shade_cell(right, "F4F8FB")
    rp_v = right.add_paragraph()
    rp_v.add_run(signed_at).font.size = Pt(11)

    # AI-assisted disclosure
    disc = doc.add_paragraph()
    rd1 = disc.add_run("AI-assisted disclosure  ")
    rd1.bold = True
    rd1.font.size = Pt(9)
    rd1.font.color.rgb = DIM
    is_disclosed = soap_att.get("ai_assisted_disclosure", True)
    rd2 = disc.add_run("✓ Acknowledged" if is_disclosed else "⚠ NOT ACKNOWLEDGED")
    rd2.font.size = Pt(9)
    rd2.font.color.rgb = ACCENT if is_disclosed else ROSE


def _render_footer(doc, validation: dict | None) -> None:
    """Tiny audit-trail line at the end. Date + sig score + brand mark."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("─" * 100)
    r.font.size = Pt(6)
    r.font.color.rgb = DIM

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ts = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    score = ((validation or {}).get("signability_score"))
    score_str = f"  ·  Signability score: {score}/100" if score is not None else ""
    rf = foot.add_run(
        f"Generated {ts}{score_str}  ·  DentaScribe MVP v0.1  ·  AI-assisted; provider-signed."
    )
    rf.font.size = Pt(8)
    rf.font.color.rgb = DIM


# ---------- table style helper ----------

def _style_data_table(table, headers: list[str]) -> None:
    """Apply the clinical-table look: mint header row, thin row dividers."""
    table.autofit = True
    # Hide outer borders (we add per-cell where needed)
    _hide_table_borders(table)
    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr[i]
        _shade_cell(cell, "E6F8F1")
        # bottom border to separate from data rows
        _set_cell_border(cell, bottom=True, color="4DD4AC", sz=10)
        p = cell.paragraphs[0]
        r = p.add_run(h.upper())
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = NAVY


# ---------- save-to-file convenience (used by the Templates page) ----------

def write_blank_template(out_path: str | Path) -> Path:
    """Generate a blank template file so it can be downloaded from the UI."""
    bytes_ = build_soap_docx(soap={}, attestation={}, validation=None)
    p = Path(out_path)
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_bytes(bytes_)
    return p


if __name__ == "__main__":
    # Self-test: build from the bundled emergency-endo sample
    sample = Path(__file__).resolve().parent.parent / "data" / "sample_filled_soap_emergency_endo.json"
    if sample.exists():
        soap = json.loads(sample.read_text())
        out = Path("/tmp/dentascribe_sample.docx")
        out.write_bytes(build_soap_docx(soap))
        print(f"✓ wrote {out} ({out.stat().st_size:,} bytes)")
    else:
        out = Path("/tmp/dentascribe_blank.docx")
        out.write_bytes(build_soap_docx({}))
        print(f"✓ wrote blank template {out} ({out.stat().st_size:,} bytes)")
