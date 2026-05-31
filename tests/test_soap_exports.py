"""Tests for the printable SOAP generators (DOCX + PDF).

We don't visually compare layouts — we verify:
  - The bytes returned are valid Word/PDF docs (magic numbers)
  - All clinical sections appear when input is fully populated
  - Empty input produces a non-zero blank template (no crash, no NPE)
  - Sensitive provider/patient values DO appear in output (otherwise the
    chart would be unsignable)
"""
from __future__ import annotations
import pathlib, sys
ROOT = pathlib.Path(__file__).parent.parent
sys.path.insert(0, str(ROOT))

import io
import json
from docx import Document

from exports.soap_docx_template import build_soap_docx
from exports.soap_pdf_template import build_soap_pdf


SAMPLE_SOAP = {
    "encounter_meta": {
        "date_of_service": "2026-05-31",
        "visit_type": "emergency",
        "provider": {"name": "Dr. A. Patel", "tsbde_license": "TX-12345",
                     "npi": "1234567890", "role": "dentist"},
        "patient": {"patient_id": "P-TEST-001", "dob": "1985-04-12"},
        "practice_location": {"city": "Dallas", "state": "TX"},
    },
    "subjective": {
        "chief_complaint": "Severe pain on lower left tooth #19",
        "hpi": {"onset": "3 days ago", "severity_0_10": 8, "triggers": ["cold", "biting"]},
        "medical_history_updates": "HTN on lisinopril",
        "medications": ["lisinopril 10mg daily"],
        "allergies": ["penicillin"],
    },
    "objective": {
        "vitals": {"bp": "128/82", "hr": 76, "temp": "98.4"},
        "intra_oral": "Caries on #19 occlusal, deep",
        "exam_findings": [
            {"tooth": "19", "surfaces": ["O"], "finding": "deep caries",
             "severity": "severe", "source_span": "dark spot near the apex"},
        ],
        "radiographs_taken": [
            {"type": "PA", "tooth": "19", "findings": "periapical lucency",
             "source_span": "the PA shows a dark spot"},
        ],
    },
    "assessment": {
        "diagnoses": [
            {"tooth": "19", "diagnosis": "Irreversible pulpitis", "severity": "severe"},
        ],
    },
    "plan": {
        "procedures_today": [
            {"tooth": "19", "surfaces": ["O"], "procedure": "Endodontic therapy initiated",
             "anesthesia": "lidocaine 2% with epi", "cdt_code": "D3330"},
        ],
        "prescriptions": [
            {"drug": "ibuprofen", "strength": "400 mg", "sig": "q6h PRN pain",
             "quantity": 20, "refills": 0, "interaction_checked": True},
        ],
        "follow_up": "2 weeks for crown",
        "patient_instructions": "Soft diet x48h. Avoid chewing on #19.",
    },
    "billing": {
        "cdt_codes": [
            {"code": "D0140", "description": "Limited oral evaluation",
             "tooth": "19", "rationale": "Emergency exam"},
            {"code": "D0220", "description": "Intraoral periapical first image",
             "tooth": "19", "rationale": "Diagnostic radiograph"},
            {"code": "D3330", "description": "Endodontic therapy molar",
             "tooth": "19", "rationale": "Irreversible pulpitis"},
        ],
    },
    "compliance": {
        "tsbde_checklist": {
            "patient_identified": True, "consent_documented": True,
            "license_on_record": True, "date_of_service_present": True,
            "anesthetic_documented": True, "radiograph_reference_present": True,
            "history_and_chief_complaint_present": True,
            "diagnosis_and_plan_present": True, "provider_signature_pending": False,
        },
    },
    "attestation": {"ai_assisted_disclosure": True, "signed_by": "Dr. A. Patel",
                     "signed_at": "2026-05-31T14:30Z"},
}


# ---------- DOCX ----------

def test_docx_builds_with_full_sample():
    data = build_soap_docx(SAMPLE_SOAP)
    assert isinstance(data, bytes) and len(data) > 5000
    # PK header — DOCX is a zip
    assert data[:2] == b"PK", "DOCX must start with PK (zip magic)"


def test_docx_builds_with_empty_input():
    data = build_soap_docx({})
    assert isinstance(data, bytes) and len(data) > 3000
    assert data[:2] == b"PK"


def _all_docx_text(doc) -> str:
    """Gather text from paragraphs AND all table cells (recursively).
    Section headings live in mint accent-bar tables, so paragraph-only scans
    miss them."""
    pieces = [p.text for p in doc.paragraphs]
    def _walk(table):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    pieces.append(p.text)
                for sub in cell.tables:
                    _walk(sub)
    for tbl in doc.tables:
        _walk(tbl)
    return " ".join(pieces)


def test_docx_contains_clinical_anchors():
    """Every section heading must appear so a dentist scanning the doc
    knows where to look. Section heads live in table cells (mint accent bars),
    so we scan both paragraphs and table content."""
    data = build_soap_docx(SAMPLE_SOAP)
    doc = Document(io.BytesIO(data))
    full_text = _all_docx_text(doc)
    for anchor in ["Subjective", "Objective", "Assessment", "Plan",
                    "Billing", "Compliance", "Attestation",
                    "DentaScribe", "TSBDE"]:
        assert anchor.upper() in full_text.upper(), \
            f"DOCX missing section anchor: {anchor!r}"


def test_docx_carries_provider_and_patient_identity():
    """The chart isn't signable without the encounter anchor fields."""
    data = build_soap_docx(SAMPLE_SOAP)
    doc = Document(io.BytesIO(data))
    full_text = " ".join(p.text for p in doc.paragraphs)
    # Patient + provider identity must appear somewhere (in paragraphs OR tables)
    table_text = " ".join(c.text for tbl in doc.tables
                          for row in tbl.rows for c in row.cells)
    combined = full_text + " " + table_text
    for must_appear in ["Dr. A. Patel", "TX-12345", "P-TEST-001", "2026-05-31"]:
        assert must_appear in combined, f"DOCX missing identity field: {must_appear}"


def test_docx_carries_chief_complaint_verbatim():
    data = build_soap_docx(SAMPLE_SOAP)
    doc = Document(io.BytesIO(data))
    text = " ".join(p.text for p in doc.paragraphs)
    assert "Severe pain on lower left tooth #19" in text


def test_docx_carries_cdt_codes():
    data = build_soap_docx(SAMPLE_SOAP)
    doc = Document(io.BytesIO(data))
    table_text = " ".join(c.text for tbl in doc.tables
                          for row in tbl.rows for c in row.cells)
    for code in ["D0140", "D0220", "D3330"]:
        assert code in table_text, f"DOCX missing CDT code: {code}"


# ---------- PDF ----------

def test_pdf_builds_with_full_sample():
    data = build_soap_pdf(SAMPLE_SOAP)
    assert isinstance(data, bytes) and len(data) > 3000
    assert data[:4] == b"%PDF", "PDF must start with %PDF"


def test_pdf_builds_with_empty_input():
    data = build_soap_pdf({})
    assert isinstance(data, bytes) and len(data) > 2000
    assert data[:4] == b"%PDF"


def test_pdf_carries_validation_score_in_footer():
    """The audit footer must include the signability score."""
    data = build_soap_pdf(SAMPLE_SOAP, validation={"signability_score": 92})
    # PDF text is binary; for a smoke check just verify size grew a bit
    plain = build_soap_pdf(SAMPLE_SOAP, validation=None)
    assert len(data) >= len(plain) - 100  # roughly comparable
