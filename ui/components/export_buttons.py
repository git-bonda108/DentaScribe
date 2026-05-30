"""Download buttons — populated JSON / DOCX / blank template."""
from __future__ import annotations
import json, io
import streamlit as st


def _docx_bytes(soap, attestation):
    try:
        from docx import Document
    except ImportError:
        return b""
    doc = Document()
    doc.add_heading("DentaScribe — SOAP Note", level=1)
    md = soap.get("metadata", {})
    doc.add_paragraph(f"Encounter: {md.get('encounter_id','—')}")
    doc.add_paragraph(f"Visit type: {md.get('visit_type','—')}")
    for section_key in ["subjective", "objective", "assessment", "plan"]:
        section = soap.get(section_key, {})
        doc.add_heading(section_key.title(), level=2)
        doc.add_paragraph(json.dumps(section, indent=2))
    if attestation:
        doc.add_heading("Attestation", level=2)
        doc.add_paragraph(attestation.get("signed_text", ""))
        doc.add_paragraph(f"Signed by: {attestation.get('provider_name')} "
                          f"(TX #{attestation.get('provider_license')})")
        doc.add_paragraph(f"Signed at: {attestation.get('signed_at')}")
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


def render_export_buttons(soap, attestation, blank_template=None) -> None:
    c1, c2, c3 = st.columns(3)
    with c1:
        if soap:
            st.download_button(
                "⬇  Populated SOAP (JSON)",
                data=json.dumps(soap, indent=2).encode("utf-8"),
                file_name="soap_populated.json", mime="application/json",
                use_container_width=True,
            )
        else:
            st.button("⬇  Populated SOAP (JSON)", disabled=True, use_container_width=True)
    with c2:
        if soap:
            data = _docx_bytes(soap, attestation)
            st.download_button(
                "⬇  Populated SOAP (DOCX)",
                data=data, file_name="soap_populated.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True, disabled=not data,
            )
        else:
            st.button("⬇  Populated SOAP (DOCX)", disabled=True, use_container_width=True)
    with c3:
        if blank_template:
            st.download_button(
                "⬇  Blank Texas SOAP Template",
                data=json.dumps(blank_template, indent=2).encode("utf-8"),
                file_name="texas_blank_soap_template.json", mime="application/json",
                use_container_width=True,
            )
        else:
            st.button("⬇  Blank Texas SOAP Template", disabled=True, use_container_width=True)
